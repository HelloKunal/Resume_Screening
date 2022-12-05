import docx
import re
from resume_rating import image_handler, ppt_handler
def get_content_as_string(document_path):
    fullText = ''
    if document_path.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
        fullText = image_handler.jpg_to_string(document_path)
    elif document_path.lower().endswith(('pptx', 'pptm', 'ppt')):
        fullText = ppt_handler.ppt_to_string(document_path).lower()
    else:
        req_doc = docx.Document(document_path)

        for para in req_doc.paragraphs:
            fullText = fullText + ' ' + para.text

    tbl = req_doc.tables
    for table in req_doc.tables:
        for row in table.rows:
            rowText = ''
            for cell in row.cells:
                rowText = rowText + ' ' + cell.text
            fullText = fullText + ' '+ rowText
    fullText = re.sub('\W+', ' ', fullText) #Select only alpha numerics
    fullText = re.sub('[^A-Za-z]+', ' ', fullText) #select only alphabet characters
    fullText = fullText.lower()

    return fullText
